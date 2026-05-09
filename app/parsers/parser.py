"""
app/parsers/parser.py

Handles PDF and DOCX extraction for both resumes and job descriptions.
Returns clean plain text, stripping formatting artifacts.
"""
from __future__ import annotations
import io
import logging
import re

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect format from filename extension and extract plain text.

    Args:
        file_bytes: Raw file bytes (from upload or disk)
        filename:   Original filename — used to determine format

    Returns:
        Cleaned plain text string
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: pdf, docx, txt")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages)
        logger.info("PDF extracted: %d pages, %d chars", len(pages), len(text))
        return _clean(text)
    except ImportError:
        # Fallback to pdfminer
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        output = io.StringIO()
        extract_text_to_fp(io.BytesIO(file_bytes), output, laparams=LAParams())
        text = output.getvalue()
        logger.info("PDF extracted via pdfminer: %d chars", len(text))
        return _clean(text)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cell text (common in resume skill tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs)
    logger.info("DOCX extracted: %d paragraphs, %d chars", len(paragraphs), len(text))
    return _clean(text)


def _clean(text: str) -> str:
    """Normalise whitespace, remove non-printable chars."""
    text = re.sub(r"[^\x09\x0a\x0d\x20-\x7e\u00a0-\uffff]", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
